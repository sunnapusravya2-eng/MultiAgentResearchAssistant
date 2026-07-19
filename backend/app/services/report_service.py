from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from fpdf import FPDF

from app.core.config import DATA_DIR
from app.core.logging import logger
from app.rag.retriever import RetrieverAgent
from app.services.llm_service import create_llm

REPORTS_DIR = DATA_DIR / "reports"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)


class ReportService:
    def __init__(self) -> None:
        self.retriever = RetrieverAgent()
        self.llm = create_llm()
        self.logger = logger

    def generate_report(self, title: str, question: str, conversation_id: str) -> dict[str, Any]:
        if not title.strip() or not question.strip():
            raise ValueError("Title and question are required.")

        context = self.retriever.build_context(question, top_k=8)
        prompt = self._build_report_prompt(title=title, question=question, context=context)

        self.logger.info("Generating report with LLM for title %s", title)
        report_body = self.llm.invoke(prompt).content

        pdf_path = self._save_pdf(title=title, content=report_body)
        docx_path = self._save_docx(title=title, content=report_body)

        return {
            "title": title,
            "question": question,
            "conversation_id": conversation_id,
            "content": report_body,
            "pdf_path": str(pdf_path),
            "docx_path": str(docx_path),
        }

    @staticmethod
    def _build_report_prompt(title: str, question: str, context: str) -> str:
        return (
            "You are a research assistant writing a formal report.\n\n"
            f"Report title: {title}\n"
            f"Research question: {question}\n\n"
            "Use the following context extracted from uploaded documents to write "
            "a well-structured, factual report. Include an introduction, key findings, "
            "and a conclusion. Do not use markdown tables. Keep formatting simple "
            "(plain paragraphs and dashes for bullet points).\n\n"
            f"Context:\n{context}\n\n"
            "Write the full report now."
        )

    @staticmethod
    def _safe_filename(title: str) -> str:
        return re.sub(r"[^a-zA-Z0-9_-]+", "_", title.strip()).strip("_") or "report"

    def _save_pdf(self, title: str, content: str) -> Path:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        effective_width = pdf.w - pdf.l_margin - pdf.r_margin

        def break_long_word(word: str, max_width: float) -> str:
            if pdf.get_string_width(word) <= max_width:
                return word
            chunks = []
            current = ""
            for ch in word:
                test = current + ch
                if pdf.get_string_width(test) > max_width * 0.95:
                    chunks.append(current)
                    current = ch
                else:
                    current = test
            if current:
                chunks.append(current)
            return " ".join(chunks)

        def sanitize_line(line: str) -> str:
            safe = line.encode("latin-1", "replace").decode("latin-1")
            words = safe.split(" ")
            words = [break_long_word(w, effective_width) if w else w for w in words]
            return " ".join(words)

        # Title
        pdf.set_font("Arial", "B", 16)
        pdf.multi_cell(effective_width, 10, sanitize_line(title))
        pdf.ln(4)
        pdf.set_font("Arial", size=12)

        for raw_line in content.split("\n"):
            if not raw_line.strip():
                pdf.ln(4)
                continue
            safe_line = sanitize_line(raw_line)
            try:
                pdf.multi_cell(effective_width, 8, safe_line)
            except Exception as exc:
                self.logger.warning("Skipping unrenderable PDF line: %s", exc)
                continue

        pdf_path = REPORTS_DIR / f"{self._safe_filename(title)}.pdf"
        pdf.output(str(pdf_path))
        return pdf_path

    def _save_docx(self, title: str, content: str) -> Path:
        doc = Document()
        doc.add_heading(title, level=1)

        for raw_line in content.split("\n"):
            line = raw_line.strip()
            if not line:
                doc.add_paragraph("")
                continue
            if line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("#"):
                heading_text = line.lstrip("#").strip()
                doc.add_heading(heading_text, level=2)
            else:
                doc.add_paragraph(line)

        docx_path = REPORTS_DIR / f"{self._safe_filename(title)}.docx"
        doc.save(str(docx_path))
        return docx_path